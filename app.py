import os
import json
import streamlit as st
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain.chains import ConversationalRetrievalChain
from langchain.prompts import PromptTemplate
from langchain.schema import Document as LangchainDocument
import re
from typing import List, Dict, Any, Optional
from difflib import SequenceMatcher
import docx
import PyPDF2
from io import BytesIO
import nltk
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from openai import OpenAI
from dotenv import load_dotenv
import shutil
from improved_docx_generation import generate_improved_word_document
import pysqlite3

__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')

# # Add this at the top of your script, before other imports
load_dotenv()

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

VECTORSTORE_DIR = "./chroma_store"
PROJECT_JSON = "./Shared/Marketing/Project Descriptions/combined_project_data.json"
WEBSITE_MD = "./Shared/website.md"

USE_OPENAI = True  # Set to True to use OpenAI, False for local Ollama
def get_openai_api_key():
    """Get OpenAI API key from environment or Streamlit secrets"""
    api_key = st.secrets["OPENAI_API_KEY"]
    if not api_key:
        try:
            api_key = os.getenv('OPENAI_API_KEY')
        except Exception:
            pass
    return api_key

OPENAI_API_KEY = get_openai_api_key()
if USE_OPENAI and not OPENAI_API_KEY:
    st.error("⚠️ OpenAI API key not found. Please set OPENAI_API_KEY in your environment variables or Streamlit secrets.")
    st.stop()

openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

st.title("K&M Project Knowledge Assistant")
st.caption("Powered by comprehensive project database and website content")

# ---- CV PROCESSING FUNCTIONS ----

def extract_text_from_docx(file_content):
    try:
        doc = docx.Document(BytesIO(file_content))
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text.append(paragraph.text.strip())
                    if cell_text:
                        row_text.append(' '.join(cell_text))
                if row_text:
                    full_text.append(' | '.join(row_text))
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return None

def extract_text_from_pdf(file_content):
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text.append(page_text.strip())
        return '\n'.join(text)
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return None

def separate_cv_sections(cv_text: str) -> tuple:
    work_undertaken_patterns = [
        r'Work\s+Undertaken\s+that\s+Best\s+Illustrates\s+Capability\s+to\s+Handle\s+the\s+Tasks\s+Assigned',
        r'Work\s+Undertaken\s+that\s+Best\s+Illustrates',
        r'WORK\s+UNDERTAKEN\s+THAT\s+BEST\s+ILLUSTRATES',
        r'Work\s+undertaken\s+that\s+best\s+illustrates',
        r'Project\s+Experience',
        r'PROJECT\s+EXPERIENCE',
        r'Key\s+Projects',
        r'Selected\s+Projects'
    ]
    project_section_start = None
    for pattern in work_undertaken_patterns:
        match = re.search(pattern, cv_text, re.IGNORECASE)
        if match:
            project_section_start = match.start()
            break
    if project_section_start is not None:
        bio_text = cv_text[:project_section_start].strip()
        projects_text = cv_text[project_section_start:].strip()
    else:
        project_patterns = [
            r'Project\s+Director[,:]',
            r'Project\s+Manager[,:]',
            r'Team\s+Leader[,:]',
            r'Lead\s+Consultant[,:]',
            r'Senior\s+\w+\s+Expert[,:]'
        ]
        matches = []
        for pattern in project_patterns:
            for match in re.finditer(pattern, cv_text, re.IGNORECASE):
                matches.append(match.start())
        if matches:
            project_section_start = min(matches)
            bio_text = cv_text[:project_section_start].strip()
            projects_text = cv_text[project_section_start:].strip()
        else:
            split_point = int(len(cv_text) * 0.6)
            bio_text = cv_text[:split_point].strip()
            projects_text = cv_text[split_point:].strip()
    return bio_text, projects_text

def parse_projects_section(projects_text: str) -> List[str]:
    project_separators = [
        r'Project\s+Director[,:]',
        r'Project\s+Manager[,:]',
        r'Team\s+Leader[,:]',
        r'Lead\s+Consultant[,:]',
        r'Senior\s+\w+\s+Expert[,:]',
        r'PPP\s+Financial\s+Expert[,:]',
        r'Energy\s+Economist[,:]',
        r'Technical\s+Manager[,:]',
        r'LNG\s+Expert[,:]',
        r'Quality\s+Control[,:]',
        r'Quality\s+Assurance[,:]'
    ]
    project_starts = []
    for pattern in project_separators:
        for match in re.finditer(pattern, projects_text, re.IGNORECASE):
            project_starts.append((match.start(), match.group()))
    project_starts.sort(key=lambda x: x[0])
    if not project_starts:
        lines = projects_text.split('\n')
        projects = []
        current_project = []
        for line in lines:
            line = line.strip()
            if line:
                if re.match(r'^[A-Z][^,]*,\s*[A-Z][^(]*\([0-9]{4}', line):
                    if current_project:
                        projects.append('\n'.join(current_project))
                        current_project = []
                current_project.append(line)
        if current_project:
            projects.append('\n'.join(current_project))
        return projects
    projects = []
    for i, (start_pos, _) in enumerate(project_starts):
        if i < len(project_starts) - 1:
            end_pos = project_starts[i + 1][0]
            project_text = projects_text[start_pos:end_pos].strip()
        else:
            project_text = projects_text[start_pos:].strip()
        if project_text:
            projects.append(project_text)
    return projects

def find_relevant_projects_with_gpt(projects_text: str, user_description: str, use_gpt: bool = True) -> List[str]:
    if not use_gpt:
        st.info("GPT processing disabled - returning empty project list")
        return []
    if not openai_client:
        st.error("OpenAI client not initialized. Please check your API key.")
        return []
    
    # Debug: Show projects text length
    st.info(f"Projects text length: {len(projects_text)} characters")
    
    # Parse projects first to get count
    individual_projects = parse_projects_section(projects_text)
    st.info(f"Parsed {len(individual_projects)} individual projects from CV")
    
    # Show first few project titles for debugging
    if individual_projects:
        st.info("First few project titles:")
        for i, proj in enumerate(individual_projects[:3]):
            title = proj.split('\n')[0][:100] + "..." if len(proj.split('\n')[0]) > 100 else proj.split('\n')[0]
            st.write(f"{i+1}. {title}")
    
    # Estimate token count (rough approximation: 1 token ≈ 4 characters)
    estimated_tokens = len(projects_text) / 4 + len(user_description) / 4 + 500  # +500 for prompt
    st.info(f"Estimated tokens: {estimated_tokens:.0f}")
    
    # If estimated tokens exceed ~50K or more than 15 projects, chunk for better processing
    if estimated_tokens > 50000 or len(individual_projects) > 15:
        st.info("Large CV detected - processing in chunks for better results...")
        return find_relevant_projects_chunked(projects_text, user_description)
    
    try:
        prompt = f"""
        You are an expert CV analyst. I will provide you with a projects section from a CV and a job description. 
        Your task is to identify ALL projects from the CV that have ANY relevance to the job requirements.
        
        IMPORTANT: Be INCLUSIVE rather than selective. If a project has even partial relevance, include it.
        Look for relevance based on:
        - Similar or related locations/regions
        - Similar or related technologies/sectors
        - Similar or related services/roles
        - Similar or related project types
        - Similar or related skills required
        - Any transferable experience

        JOB DESCRIPTION:
        {user_description}

        CV PROJECTS SECTION:
        {projects_text}

        CRITICAL FORMATTING INSTRUCTIONS:
        1. Carefully read through ALL the projects in the CV
        2. For each project that has ANY relevance, copy the EXACT project description
        3. Put each relevant project between these exact markers:
           ---PROJECT_SEPARATOR---
           [FULL PROJECT DESCRIPTION HERE]
           ---PROJECT_SEPARATOR---
        4. At the very end, add: "TOTAL RELEVANT PROJECTS FOUND: X"
        
        EXAMPLE FORMAT:
        ---PROJECT_SEPARATOR---
        Project Manager, Example Project, Location (Year): Full project description here...
        ---PROJECT_SEPARATOR---
        ---PROJECT_SEPARATOR---
        Another Project Manager, Second Project, Location (Year): Another full description...
        ---PROJECT_SEPARATOR---
        TOTAL RELEVANT PROJECTS FOUND: 2

        Be generous in your interpretation of relevance. Include projects with even partial relevance.
        If absolutely no projects are relevant, return only: "NO_RELEVANT_PROJECTS"
        """
        
        st.info("Sending single request to GPT for project matching...")
        response = openai_client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are an expert CV analyst. Your goal is to be INCLUSIVE and find ALL projects with ANY relevance to the job requirements."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=20000  # Increased from 16000 to handle more projects
        )
        result = response.choices[0].message.content.strip()
        
        st.info(f"GPT response length: {len(result)} characters")
        
        # Show the summary line if present
        if "TOTAL RELEVANT PROJECTS FOUND:" in result:
            summary_line = [line for line in result.split('\n') if "TOTAL RELEVANT PROJECTS FOUND:" in line]
            if summary_line:
                st.write(f"GPT Summary: {summary_line[0]}")
        
        if result == "NO_RELEVANT_PROJECTS":
            st.warning("GPT found no relevant projects")
            return []
        
        # Remove the summary line before processing
        result_clean = result.split("TOTAL RELEVANT PROJECTS FOUND:")[0].strip()
        
        # Count claimed vs actual projects
        claimed_count = 0
        if "TOTAL RELEVANT PROJECTS FOUND:" in result:
            try:
                claimed_count = int(result.split("TOTAL RELEVANT PROJECTS FOUND:")[1].strip().split()[0])
            except:
                claimed_count = 0
        
        relevant_projects = [proj.strip() for proj in result_clean.split("---PROJECT_SEPARATOR---") if proj.strip()]
        actual_count = len(relevant_projects)
        
        # Validation: Check if GPT claimed more projects than it actually provided
        if claimed_count > actual_count and claimed_count > 1:
            st.error(f"⚠️ GPT FORMATTING ERROR: Claimed {claimed_count} projects but only provided {actual_count}")
            st.error("This indicates GPT didn't use the separator format correctly")
            st.write("Raw GPT response:")
            st.code(result[:1000] + "..." if len(result) > 1000 else result)
        
        st.success(f"GPT found {actual_count} relevant projects (claimed: {claimed_count})")
        
        # Show titles of found projects for debugging
        if relevant_projects:
            st.write("Found project titles:")
            for i, proj in enumerate(relevant_projects):
                title = proj.split('\n')[0][:80] + "..." if len(proj.split('\n')[0]) > 80 else proj.split('\n')[0]
                st.write(f"  {i+1}. {title}")
        
        return relevant_projects
    except Exception as e:
        st.error(f"Error calling OpenAI API for project matching: {e}")
        return []

def find_relevant_projects_chunked(projects_text: str, user_description: str) -> List[str]:
    """
    Process large CVs in chunks to handle token limits
    """
    try:
        # Parse individual projects first
        individual_projects = parse_projects_section(projects_text)
        
        if not individual_projects:
            st.warning("No projects found in chunked processing")
            return []
        
        st.info(f"Chunked processing: Found {len(individual_projects)} total projects")
        
        # Process projects in chunks of ~10-12 projects each (smaller chunks for better processing)
        chunk_size = 10
        all_relevant_projects = []
        
        for i in range(0, len(individual_projects), chunk_size):
            chunk = individual_projects[i:i + chunk_size]
            chunk_text = "\n\n".join(chunk)
            
            st.info(f"Processing chunk {i//chunk_size + 1} of {(len(individual_projects) + chunk_size - 1)//chunk_size} ({len(chunk)} projects)...")
            
            # Show first project in chunk for debugging
            if chunk:
                first_title = chunk[0].split('\n')[0][:100] + "..." if len(chunk[0].split('\n')[0]) > 100 else chunk[0].split('\n')[0]
                st.write(f"First project in chunk: {first_title}")
                
                # Show all project titles in this chunk for debugging
                with st.expander(f"📋 All {len(chunk)} projects in chunk {i//chunk_size + 1}"):
                    for j, proj in enumerate(chunk):
                        title = proj.split('\n')[0][:120] + "..." if len(proj.split('\n')[0]) > 120 else proj.split('\n')[0]
                        st.write(f"{j+1}. {title}")
            
            prompt = f"""
            You are an expert CV analyst. I will provide you with a subset of projects from a CV and a job description. 
            Your task is to identify ALL projects from this subset that have ANY relevance to the job requirements.
            
            IMPORTANT: Be INCLUSIVE rather than selective. If a project has even partial relevance, include it.
            Look for relevance based on:
            - Similar or related locations/regions
            - Similar or related technologies/sectors  
            - Similar or related services/roles
            - Similar or related project types
            - Similar or related skills required
            - Any transferable experience

            JOB DESCRIPTION:
            {user_description}

            CV PROJECTS SUBSET ({len(chunk)} projects):
            {chunk_text}

            CRITICAL FORMATTING INSTRUCTIONS:
            1. Go through EACH of the {len(chunk)} projects in this subset
            2. For each project that has ANY relevance, copy the EXACT project description
            3. Put each relevant project between these exact markers:
               ---PROJECT_SEPARATOR---
               [FULL PROJECT DESCRIPTION HERE]
               ---PROJECT_SEPARATOR---
            4. At the very end, add: "TOTAL RELEVANT PROJECTS FOUND: X"
            
            EXAMPLE FORMAT:
            ---PROJECT_SEPARATOR---
            Project Manager, Example Project, Location (Year): Full project description here...
            ---PROJECT_SEPARATOR---
            ---PROJECT_SEPARATOR---
            Another Project Manager, Second Project, Location (Year): Another full description...
            ---PROJECT_SEPARATOR---
            TOTAL RELEVANT PROJECTS FOUND: 2

            Be generous in your interpretation of relevance. Include projects with even partial relevance.
            If absolutely no projects are relevant, return only: "NO_RELEVANT_PROJECTS"
            """
            
            response = openai_client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are an expert CV analyst. Your goal is to be INCLUSIVE and find ALL projects with ANY relevance to the job requirements."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=20000  # Increased to handle more projects
            )
            
            result = response.choices[0].message.content.strip()
            st.info(f"Chunk {i//chunk_size + 1} response length: {len(result)} characters")
            
            # Show the summary line if present
            if "TOTAL RELEVANT PROJECTS FOUND:" in result:
                summary_line = [line for line in result.split('\n') if "TOTAL RELEVANT PROJECTS FOUND:" in line]
                if summary_line:
                    st.write(f"GPT Summary: {summary_line[0]}")
            
            if result != "NO_RELEVANT_PROJECTS":
                # Remove the summary line before processing
                result_clean = result.split("TOTAL RELEVANT PROJECTS FOUND:")[0].strip()
                
                # Count claimed vs actual projects
                claimed_count = 0
                if "TOTAL RELEVANT PROJECTS FOUND:" in result:
                    try:
                        claimed_count = int(result.split("TOTAL RELEVANT PROJECTS FOUND:")[1].strip().split()[0])
                    except:
                        claimed_count = 0
                
                chunk_relevant_projects = [proj.strip() for proj in result_clean.split("---PROJECT_SEPARATOR---") if proj.strip()]
                actual_count = len(chunk_relevant_projects)
                
                # Validation: Check if GPT claimed more projects than it actually provided
                if claimed_count > actual_count and claimed_count > 1:
                    st.error(f"⚠️ GPT FORMATTING ERROR: Claimed {claimed_count} projects but only provided {actual_count}")
                    st.error("This indicates GPT didn't use the separator format correctly")
                    st.write("Raw GPT response:")
                    st.code(result[:1000] + "..." if len(result) > 1000 else result)
                
                st.success(f"Chunk {i//chunk_size + 1} found {actual_count} relevant projects out of {len(chunk)} total (claimed: {claimed_count})")
                all_relevant_projects.extend(chunk_relevant_projects)
                
                # Show titles of found projects for debugging
                if chunk_relevant_projects:
                    st.write("Found project titles:")
                    for j, proj in enumerate(chunk_relevant_projects):
                        title = proj.split('\n')[0][:80] + "..." if len(proj.split('\n')[0]) > 80 else proj.split('\n')[0]
                        st.write(f"  {j+1}. {title}")
            else:
                st.info(f"Chunk {i//chunk_size + 1} found no relevant projects")
        
        st.success(f"Total: Found {len(all_relevant_projects)} relevant projects from {len(individual_projects)} total projects")
        return all_relevant_projects
        
    except Exception as e:
        st.error(f"Error processing chunked projects: {e}")
        return []

def generate_qualification_paragraphs_with_gpt(bio_text: str, relevant_projects: List[str], user_description: str, use_gpt: bool = True) -> str:
    if not use_gpt:
        st.info("GPT processing disabled - returning basic qualification text")
        return "Qualification generation disabled. Please enable GPT processing to generate tailored qualification paragraphs."
    if not openai_client:
        st.error("OpenAI client not initialized. Please check your API key.")
        return "Unable to generate qualifications - OpenAI API not available."
    try:
        projects_combined = "\n\n".join(relevant_projects) if relevant_projects else "No directly relevant projects identified."
        prompt = f"""
        You are an expert proposal writer. I will provide you with:
        1. Bio section from a CV
        2. Relevant project experience 
        3. Job/work description
        4. Themes of the proposal

        Your task is to write a four-paragraph professional bio for this individual. Use confident, professional language without exaggeration or flattery.

        BIO SECTION:
        {bio_text}

        RELEVANT PROJECT EXPERIENCE:
        {projects_combined}

        JOB/WORK DESCRIPTION:
        {user_description}
        
        Themes of the proposal:
        {st.session_state['extracted_themes']}
        
        The four paragraphs should be structured as follows:
        1. Paragraph 1 should introduce the individual, summarizing their current role, core areas of expertise, experience with relevant power generation technologies (e.g., LNG, renewables, storage), and countries or regions where they have worked that are relevant to the proposal.
        2. Paragraph 2 should focus on a specific theme that aligns with the needs of the proposal. After paragraph 2, insert a section titled ===Relevant Project Experience for Theme 2=== and then output a table as follows: For each of exactly 2-3 relevant project experiences (copied verbatim from the RELEVANT PROJECT EXPERIENCE section above) that reinforce the theme stated in paragraph 2, output a line starting with >>>PROJECT_START<<< and ending with >>>PROJECT_END<<<, with the project description in between. Each project should be clearly separated and include the exact position, project name, location, year, and a short description as in the CV.
        3. Paragraph 3 should focus on another specific theme. After paragraph 3, insert a section titled ===Relevant Project Experience for Theme 3=== and then output a table as follows: For each of exactly 2-3 relevant project experiences (copied verbatim from the RELEVANT PROJECT EXPERIENCE section above) that reinforce the theme stated in paragraph 3, output a line starting with >>>PROJECT_START<<< and ending with >>>PROJECT_END<<<, with the project description in between. Each project should be clearly separated and include the exact position, project name, location, year, and a short description as in the CV (if the tense in the project description is not past tense, change it to past tense, but keep the exact project description as in the CV).
        4. Paragraph 4 should conclude the bio by summarizing how the individual's experience aligns with the proposed assignment. It should also include a brief statement of their academic background, including degrees earned, fields of study, and the institutions attended. Ensure the writing is cohesive, clear, and appropriate for inclusion in a technical or commercial proposal.

        Write in third person (he/she) and make it compelling but factual.
        """
        response = openai_client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are an expert proposal writer specializing in highlighting candidate qualifications for infrastructure and consulting projects."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=2000  # Increased from 500 to handle longer qualifications
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error calling OpenAI API for qualification generation: {e}")
        return f"Unable to generate qualifications due to API error: {str(e)}"

def get_project_title(proj):
    """Extract project title from project description (everything before first colon or first line)"""
    # Use text before first colon or first line as title
    if ':' in proj:
        return proj.split(':')[0].strip()
    else:
        return proj.split('\n')[0].strip()

# ---- MOVE: Button Mode Selection to sidebar top ----
if 'mode' not in st.session_state:
    st.session_state['mode'] = 'cv'  # Default to search

# Initialize separate chat histories for each tab
if 'chat_history_general' not in st.session_state:
    st.session_state['chat_history_general'] = []

if 'chat_history_search' not in st.session_state:
    st.session_state['chat_history_search'] = []

if 'chat_history_cv' not in st.session_state:
    st.session_state['chat_history_cv'] = []

# Keep backward compatibility
if 'chat_history' not in st.session_state:
    st.session_state['chat_history'] = []
    

def get_current_chat_history():
    """Get chat history for current mode"""
    mode = st.session_state.get('mode', 'search')
    if mode == 'general':
        return st.session_state['chat_history_general']
    elif mode == 'search':
        return st.session_state['chat_history_search']
    elif mode == 'cv':
        return st.session_state['chat_history_cv']
    else:
        return st.session_state['chat_history']

def add_to_current_chat_history(user_msg, bot_msg):
    """Add message to current mode's chat history"""
    mode = st.session_state.get('mode', 'search')
    if mode == 'general':
        st.session_state['chat_history_general'].append((user_msg, bot_msg))
    elif mode == 'search':
        st.session_state['chat_history_search'].append((user_msg, bot_msg))
    elif mode == 'cv':
        st.session_state['chat_history_cv'].append((user_msg, bot_msg))
    else:
        st.session_state['chat_history'].append((user_msg, bot_msg))



with st.sidebar:
    st.subheader("Mode Selection")
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("General", key="general_btn", 
                    type="primary" if st.session_state['mode'] == 'general' else "secondary",
                    use_container_width=True):
            st.session_state['mode'] = 'general'
            st.rerun()
    with col2:
        if st.button("Search", key="search_btn", 
                    type="primary" if st.session_state['mode'] == 'search' else "secondary",
                    use_container_width=True):
            st.session_state['mode'] = 'search'
            st.rerun()
    with col3:
        if st.button("CV", key="cv_btn", 
                    type="primary" if st.session_state['mode'] == 'cv' else "secondary",
                    use_container_width=True):
            st.session_state['mode'] = 'cv'
            st.rerun()
    st.caption(f"**Current Mode:** {st.session_state['mode'].title()}")
    st.divider()

# ---- CV MODE UI ----
if st.session_state['mode'] == 'cv':
    st.subheader("📋 AI-Powered CV Analysis & Proposal Generation")
    
    with st.expander("ℹ️ About AI-Powered CV Analysis"):
        st.markdown("""
        **This CV analysis uses OpenAI GPT-4.1-mini to:**

        1. **Intelligent Project Matching**: AI reads through all projects in the CV and identifies those most relevant to your work description

        2. **Exact Text Extraction**: Relevant project descriptions are copied exactly from the CV (no summarization or modification)

        3. **Smart Qualification Writing**: AI combines the candidate's bio and relevant projects to write compelling qualification paragraphs
        """)
        
    st.info("Upload a CV (DOCX or PDF) and describe the work requirements. AI will identify relevant projects and generate key qualification paragraphs.")

    if not OPENAI_API_KEY:
        st.error("⚠️ OpenAI API key required for CV analysis. Please set OPENAI_API_KEY in your environment variables or Streamlit secrets.")
        st.stop()

    uploaded_file = st.file_uploader(
        "Upload CV File",
        type=['docx', 'pdf'],
        help="Upload a DOCX or PDF file containing the CV to analyze"
    )

    work_description = st.text_area(
        "Describe the Work/Project Requirements",
        placeholder="e.g., 'We need a senior consultant for feasibility study of a 100MW solar project in Kenya, including technical assessment, financial modeling, and regulatory analysis.'",
        height=100,
        help="Provide a detailed description of the work requirements. AI will match this against the CV projects."
    )

    # NEW: Prompt revision textbox
    default_qual_prompt = """You are an expert proposal writer. I will provide you with:
1. Bio section from a CV
2. Relevant project experience
3. Job/work description
4. Themes of the proposal

Your task is to write a four-paragraph professional bio for this individual. Use confident, professional language without exaggeration or flattery.

BIO SECTION:
{{bio_text}}

RELEVANT PROJECT EXPERIENCE:
{{relevant_projects}}

JOB/WORK DESCRIPTION:
{{user_description}}

Themes of the proposal:
{{themes}}

The four paragraphs should be structured as follows:
1. Paragraph 1 should introduce the individual, summarizing their current role, core areas of expertise, experience with relevant power generation technologies (e.g., LNG, renewables, storage), and countries or regions where they have worked that are relevant to the proposal.
2. Paragraph 2 should focus on a specific theme that aligns with the needs of the proposal. It must begin with a topic sentence stating the theme, followed by summarizing the highlighting two to three specific project examples from the CV that illustrate the individual's experience in that area.
3. Paragraph 3 should focus on another specific theme. It must also begin with a topic sentence stating the theme, followed by summarizing the highlighting two to three specific project examples from the CV that illustrate the individual's experience in that area.
4. Paragraph 4 should conclude the bio by summarizing how the individual's experience aligns with the proposed assignment. It should also include a brief statement of their academic background, including degrees earned, fields of study, and the institutions attended. Ensure the writing is cohesive, clear, and appropriate for inclusion in a technical or commercial proposal.

Write in third person (he/she) and make it compelling but factual. Do NOT include any project tables, project markers, or special formatting. Only output the four paragraphs as plain text, each separated by a blank line.
"""
    qual_prompt = st.text_area(
        "Revise the Answer Structuring Prompt (Optional)",
        value=default_qual_prompt,
        height=180,
        help="You can revise how the answer is structured. The default is a strong proposal-style prompt."
    )

    # NEW: Extract themes button
    extract_themes_btn = st.button("🔍 Extract Themes from Requirements", type="secondary", use_container_width=True, key="extract_themes_btn")

    # NEW: Placeholder for themes extraction and revision
    # Initialize themes in session state if not exists
    if 'extracted_themes' not in st.session_state:
        st.session_state['extracted_themes'] = "(Click 'Extract Themes' above to get themes from your requirements)"
    
    themes_text = st.text_area(
        "Themes to Highlight (Edit before generating CV)",
        value=st.session_state['extracted_themes'],
        height=100,
        key="themes_textbox",
        help="AI will suggest themes to highlight from the project experience. You can edit these before generating the final answer."
    )

    # --- NEW: Manual Theme Inputs ---
    st.markdown("**Enter Three Themes to Use in the CV (these will be used for generation):**")
    theme1 = st.text_input("Theme 1", key="manual_theme_1")
    theme2 = st.text_input("Theme 2", key="manual_theme_2")
    theme3 = st.text_input("Theme 3", key="manual_theme_3")

    # --- Remove editable prompt box, use fixed prompt below ---
    five_paragraph_prompt = '''You are an expert proposal writer. I will provide you with:
1. Bio section from a CV
2. Comprehensive project experience
3. Job/work description
4. Themes of the proposal

Your task is to write a five-paragraph professional bio for this individual. Use confident, professional language without exaggeration or flattery.

BIO SECTION:
{{bio_text}}

RELEVANT PROJECT EXPERIENCE:
{{relevant_projects}}

JOB/WORK DESCRIPTION:
{{user_description}}

Themes of the proposal:
{{themes}}

The five paragraphs should be structured as follows:

1. Paragraph 1 should briefly  introduce the individual, summarizing their experience in the relevant themes to the proposal (e.g., relevant power generation technologies, countries or regions where they have worked that are relevant to the proposal).
2. Paragraph 2 should focus on a specific theme (identified below) that aligns with the needs of the proposal. It should start with a topic sentence presenting the theme, followed by a short description of 2 - 3 of the most relevant projects supporting that theme.
3. Paragraph 3 should focus on a specific theme (identified below) that aligns with the needs of the proposal. It should start with a topic sentence presenting the theme, followed by a short description of 2 - 3 of the most relevant projects supporting that theme.
4. Paragraph 4 should focus on a specific theme (identified below) that aligns with the needs of the proposal. It should start with a topic sentence presenting the theme, followed by a short description of 2 - 3 of the most relevant projects supporting that theme.
5. Paragraph 5 is a 2 - 3 sentence statement of their academic background, including degrees earned, fields of study, and the institutions attended.

Ensure the writing is cohesive, clear, and appropriate for inclusion in a technical or commercial proposal. Write in third person (he/she) and make it compelling but factual.'''

    # NEW: Generate CV button (will be enabled after themes are available)
    generate_cv_btn = st.button("Generate CV", type="primary", use_container_width=True, key="generate_cv_btn")

    show_sections = False
    use_gpt_processing = True

    show_advanced_settings = False 

    if show_advanced_settings:
        with st.expander("🔧 Advanced Settings"):
            st.info("CV analysis uses OpenAI GPT-4.1-mini for intelligent project matching and qualification generation.")
            show_sections = st.checkbox(
                "Show CV Section Breakdown", 
                value=False,
                help="Display how the CV was separated into bio and projects sections"
            )
            use_gpt_processing = st.checkbox(
                "Enable GPT Processing", 
                value=True,
                help="Enable AI processing with GPT-4.1-mini. If disabled, will only show document sections without AI analysis."
            )

    # NEW: Extract themes from requirements logic
    if extract_themes_btn and work_description:
        with st.spinner("🔍 Extracting themes from requirements..."):
            def extract_themes_from_requirements(work_description, use_gpt=True):
                if not use_gpt:
                    return []
                if not openai_client:
                    st.error("OpenAI client not initialized. Please check your API key.")
                    return []
                try:
                    # Use both work_description and the CV text (bio_text + projects_text)
                    bio_text, projects_text = separate_cv_sections(cv_text) if 'cv_text' in locals() else ("", "")
                    prompt = f"""
You are an expert CV analyst. I will provide you with a job/work description and a candidate's CV (bio and project experience).

Your task is to identify the main themes, skills, sectors, technologies, and types of experience that are most relevant for this position, but ONLY include themes that are clearly relevant to BOTH the job/work description AND the candidate's actual experience as found in the CV.

Be as comprehensive as possible: include any theme that is reasonably supported by both the job description and the candidate's CV, even if the evidence is not overwhelming. If in doubt, include the theme as long as there is some support in both sources.

When extracting themes, ensure you consider and include (if relevant):
– Countries or regions with relevant work experience
– Types of assignments (e.g., due diligence, feasibility studies, PPP structuring, financial analysis)
– Technologies (e.g., gas-fired power, offshore wind, battery storage)
– Power market experience (e.g., small island systems, emerging markets, regulated vs. liberalized markets)
– Cross-cutting skills (e.g., stakeholder engagement, regulatory support, economic modeling)

JOB/WORK DESCRIPTION:
{work_description}

CANDIDATE CV (BIO + PROJECTS):
{bio_text}\n{projects_text}

Please return a concise comma-separated list of themes (e.g., 'solar power, feasibility studies, project management, Africa, financial modeling, regulatory analysis').
Focus on specific technical skills, sectors, technologies, and experience types that are valuable for this role AND are supported by the candidate's CV. Do not include generic words like 'project' or 'experience'.
Only include a theme if it is supported by BOTH the job description and the CV.
"""
                    response = openai_client.chat.completions.create(
                        model="gpt-4.1-mini",
                        messages=[
                            {"role": "system", "content": "You are an expert CV analyst specializing in extracting key themes from job requirements and candidate CVs."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.1,
                        max_tokens=800  # Increased from 200 to handle more comprehensive themes
                    )
                    result = response.choices[0].message.content.strip()
                    # Split by comma and clean
                    themes = [t.strip() for t in result.split(',') if t.strip()]
                    return themes
                except Exception as e:
                    st.error(f"Error calling OpenAI API for theme extraction: {e}")
                    return []

            themes = extract_themes_from_requirements(work_description, use_gpt_processing)
            if themes:
                st.session_state['extracted_themes'] = ', '.join(themes)
                st.success(f"✅ Extracted {len(themes)} themes from requirements!")
                st.rerun()
            else:
                st.error("No themes could be extracted. Please check your requirements description or try again.")

    if uploaded_file and work_description:
        # This section is now handled by the Extract Themes and Generate CV buttons
        pass

    # --- UPDATED: Generate CV button logic to use manual themes and fixed prompt ---
    if generate_cv_btn and uploaded_file and work_description and theme1 and theme2 and theme3:
        user_themes = ', '.join([theme1, theme2, theme3])
        file_content = uploaded_file.read()
        if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            cv_text = extract_text_from_docx(file_content)
        elif uploaded_file.type == "application/pdf":
            cv_text = extract_text_from_pdf(file_content)
        else:
            st.error("Unsupported file type")
            st.stop()
        if not cv_text:
            st.error("Could not extract text from the uploaded file")
            st.stop()
        bio_text, projects_text = separate_cv_sections(cv_text)
        individual_projects = parse_projects_section(projects_text)
        relevant_projects = find_relevant_projects_with_gpt(projects_text, work_description, use_gpt_processing)

        # --- Deduplicate by project title ---
        st.info(f"Before deduplication: {len(relevant_projects)} projects")
        unique_projects = {}
        duplicates_found = []
        for proj in relevant_projects:
            title = get_project_title(proj)
            if title not in unique_projects:
                unique_projects[title] = proj
            else:
                duplicates_found.append(f"Duplicate title: '{title}'")
        
        if duplicates_found:
            st.warning(f"Found {len(duplicates_found)} duplicate project titles:")
            for dup in duplicates_found:
                st.write(f"  - {dup}")
        
        deduped_projects = list(unique_projects.values())
        st.info(f"After deduplication: {len(deduped_projects)} projects")
        
        # Show titles of final deduped projects
        if deduped_projects:
            st.write("Final deduplicated project titles:")
            for i, proj in enumerate(deduped_projects):
                title = get_project_title(proj)[:100] + "..." if len(get_project_title(proj)) > 100 else get_project_title(proj)
                st.write(f"  {i+1}. {title}")

        # --- Tailor each project using GPT, one at a time (or in small batches if needed) ---
        def tailor_project_with_gpt(original_project, job_description):
            # Extract the original title to preserve it
            original_title = get_project_title(original_project)
            
            tailor_prompt = f"""
You are an expert proposal writer. Here is a project description from a CV and a job/work description. 
Your task is to tailor the project description to the job/work description, keeping as much of the original text and detail as possible, 
but making it directly relevant to the job/work description. The tailored description should be between 150 and 200 words. 
Do NOT invent facts, but you may rephrase, reorganize, or clarify as needed for clarity and relevance.

CRITICAL REQUIREMENT: You MUST preserve the exact project title (everything before the colon ":"). The title should remain exactly as: "{original_title}"

JOB/WORK DESCRIPTION:
{job_description}

ORIGINAL PROJECT DESCRIPTION:
{original_project}

Return only the tailored project description with the preserved title, no commentary or extra text.
"""
            try:
                response = openai_client.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {"role": "system", "content": "You are an expert proposal writer specializing in tailoring project descriptions for CVs."},
                        {"role": "user", "content": tailor_prompt}
                    ],
                    temperature=0.2,
                    max_tokens=1200
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                return f"[Error tailoring project: {str(e)}]"

        # Only tailor for the Word document, not for website display
        tailored_projects = []
        for proj in deduped_projects:
            tailored = tailor_project_with_gpt(proj, work_description)
            # Ensure the tailored project is not empty and not hallucinated (basic check: must contain some of the original text)
            if proj[:30].strip() in tailored or proj.split(':')[0].strip() in tailored:
                tailored_projects.append(tailored)
            else:
                # Fallback: use the original if GPT output is not faithful
                tailored_projects.append(proj)

        def generate_qualification_paragraphs_manual_themes(bio_text, relevant_projects, user_description, user_themes, use_gpt=True):
            if not use_gpt:
                return "Qualification generation disabled. Please enable GPT processing to generate tailored qualification paragraphs."
            if not openai_client:
                return "Unable to generate qualifications - OpenAI API not available."
            try:
                projects_combined = "\n\n".join(relevant_projects) if relevant_projects else "No directly relevant projects identified."
                prompt = (
                    five_paragraph_prompt
                    .replace("{{bio_text}}", bio_text)
                    .replace("{{relevant_projects}}", projects_combined)
                    .replace("{{user_description}}", user_description)
                    .replace("{{themes}}", user_themes)
                )
                response = openai_client.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {"role": "system", "content": "You are an expert proposal writer specializing in highlighting candidate qualifications for infrastructure and consulting projects."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.2,
                    max_tokens=8000
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                return f"Unable to generate qualifications due to API error: {str(e)}"
        qualifications = generate_qualification_paragraphs_manual_themes(bio_text, tailored_projects, work_description, user_themes, use_gpt_processing)
        # For website display, use the original (deduped) relevant projects, not tailored
        st.session_state['cv_qualifications'] = qualifications
        st.session_state['cv_relevant_projects'] = deduped_projects
        st.session_state['cv_tailored_projects'] = tailored_projects  # Save tailored for Word only
        st.session_state['cv_uploaded_file_bytes'] = file_content  # Save file for second download button
        st.session_state['cv_generated'] = True
        st.success("✅ AI Analysis Complete!")
        st.rerun()

    # --- Display Generated CV Results ---
    if st.session_state.get('cv_generated', False):
        st.markdown("---")
        st.markdown("## 📄 Generated CV")
        
        # Re-display the results
        st.markdown("### 🎯 Key Qualifications")
        st.text_area("Key Qualification Paragraphs", st.session_state.get('cv_qualifications', ''), height=220, disabled=False, key="qualifications_display")
        
        # Show only the original (untailored) relevant projects on the website
        relevant_projects = st.session_state.get('cv_relevant_projects', [])
        if relevant_projects:
            st.markdown(f"### 📁 Relevant Project Experience")
            for i, project in enumerate(relevant_projects, 1):
                st.text_area(f"Project {i}", project, height=200, disabled=False, key=f"relevant_project_{i}_display")
        else:
            st.info("No directly relevant projects identified.")
        
        # --- IMPROVED: Generate Word document with tailored projects ---
        if uploaded_file and uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            qualifications = st.session_state.get('cv_qualifications', '')
            # Use tailored projects for the Word document
            tailored_projects = st.session_state.get('cv_tailored_projects', [])
            
            # Validate that each tailored project title matches an original project title
            original_projects = st.session_state.get('cv_relevant_projects', [])
            original_titles = [get_project_title(proj) for proj in original_projects]
            
            validated_tailored_projects = []
            excluded_projects = []
            
            for tailored_proj in tailored_projects:
                tailored_title = get_project_title(tailored_proj)
                if tailored_title in original_titles:
                    validated_tailored_projects.append(tailored_proj)
                else:
                    excluded_projects.append(tailored_title)
            
            if excluded_projects:
                st.warning(f"Excluded {len(excluded_projects)} tailored projects from Word document due to title mismatch:")
                for excluded_title in excluded_projects:
                    st.write(f"  - {excluded_title}")
            
            st.info(f"Using {len(validated_tailored_projects)} validated projects for Word document")
            
            try:
                buffer = generate_improved_word_document(uploaded_file, qualifications, validated_tailored_projects)
                st.download_button(
                    label="📄 Download as Word Document",
                    data=buffer,
                    file_name="KM_CV_Qualifications.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download the CV with GPT-generated qualifications and projects, maintaining original formatting"
                )
                # --- NEW: Download button for original project descriptions ---
                orig_file_bytes = st.session_state.get('cv_uploaded_file_bytes', None)
                if orig_file_bytes:
                    from io import BytesIO
                    orig_file = BytesIO(orig_file_bytes)
                    orig_buffer = generate_improved_word_document(orig_file, qualifications, relevant_projects)
                    st.download_button(
                        label="📄 Download as Word Document (Original Project Descriptions)",
                        data=orig_buffer,
                        file_name="KM_CV_Qualifications_ORIG_PROJECTS.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Download the CV with original project descriptions (no tailoring), maintaining original formatting"
                    )
            except Exception as e:
                st.error(f"Error generating Word document: {str(e)}")
                st.info("Please try uploading the file again or contact support if the issue persists.")
        else:
            st.info("📄 Word document download is only supported for DOCX uploads.")
    elif uploaded_file:
        st.info("👆 Please describe the work requirements to analyze the CV")
    elif work_description:
        st.info("👆 Please upload a CV file to analyze")


# ---- EXISTING CODE CONTINUES (All the original project search functionality) ----

# Complete list of all countries from your images
KNOWN_COUNTRIES = {
    # A
    'argentina': ['argentina', 'argentinian'],
    'armenia': ['armenia', 'armenian'],
    'aruba': ['aruba'],
    'austria': ['austria', 'austrian'],
    'azerbaijan': ['azerbaijan'],
    
    # B
    'bahamas': ['bahamas', 'the bahamas'],
    'bangladesh': ['bangladesh'],
    'barbados': ['barbados'],
    'belize': ['belize'],
    'bolivia': ['bolivia', 'bolivian'],
    'bosnia herzegovina': ['bosnia herzegovina', 'bosnia', 'herzegovina'],
    'botswana': ['botswana'],
    'brazil': ['brazil', 'brazilian'],
    'bulgaria': ['bulgaria', 'bulgarian'],
    
    # C
    'canada': ['canada', 'canadian'],
    'caribbean': ['caribbean'],
    'cayman islands': ['cayman islands', 'cayman'],
    'chile': ['chile', 'chilean'],
    'china': ['china', 'chinese'],
    'colombia': ['colombia', 'colombian'],
    'costa rica': ['costa rica', 'costa rican'],
    'côte d\'ivoire': ['côte d\'ivoire', 'cote d\'ivoire', 'ivory coast'],
    'curacao': ['curacao', 'curaçao'],
    'czech republic': ['czech republic', 'czech'],
    
    # D
    'dominica': ['dominica'],
    'dominican republic': ['dominican republic', 'dominican', 'dr'],
    
    # E
    'ecuador': ['ecuador', 'ecuadorian'],
    'egypt': ['egypt', 'egyptian'],
    'el salvador': ['el salvador', 'salvador', 'salvadoran'],
    'estonia': ['estonia', 'estonian'],
    
    # G
    'gabon': ['gabon'],
    'germany': ['germany', 'german'],
    'ghana': ['ghana', 'ghanaian'],
    'global': ['global', 'worldwide', 'international'],
    'grenada': ['grenada'],
    'guam': ['guam'],
    'guatemala': ['guatemala', 'guatemalan'],
    'guinea': ['guinea'],
    'guinea-bissau': ['guinea-bissau', 'guinea bissau'],
    'guyana': ['guyana'],
    
    # H
    'haiti': ['haiti', 'haitian'],
    'honduras': ['honduras', 'honduran'],
    'hungary': ['hungary', 'hungarian'],
    
    # I
    'india': ['india', 'indian'],
    'indonesia': ['indonesia', 'indonesian'],
    
    # J
    'jamaica': ['jamaica', 'jamaican'],
    'jordan': ['jordan', 'jordanian'],
    
    # K
    'kenya': ['kenya', 'kenyan'],
    'korea': ['korea', 'korean', 'south korea'],
    'kyrgyz republic': ['kyrgyz republic', 'kyrgyzstan'],
    
    # L
    'laos': ['laos', 'lao'],
    'lebanon': ['lebanon', 'lebanese'],
    'liberia': ['liberia', 'liberian'],
    'lithuania': ['lithuania', 'lithuanian'],
    
    # M
    'madagascar': ['madagascar'],
    'malawi': ['malawi'],
    'malaysia': ['malaysia', 'malaysian'],
    'maldives': ['maldives'],
    'mauritania': ['mauritania', 'mauritanian'],
    'mauritius': ['mauritius'],
    'mexico': ['mexico', 'mexican'],
    'mongolia': ['mongolia', 'mongolian'],
    'morocco': ['morocco', 'moroccan'],
    
    # N
    'namibia': ['namibia', 'namibian'],
    'nepal': ['nepal', 'nepalese'],
    'netherlands': ['netherlands', 'dutch'],
    'nigeria': ['nigeria', 'nigerian'],
    
    # O
    'oman': ['oman'],
    
    # P
    'pakistan': ['pakistan', 'pakistani'],
    'panama': ['panama', 'panamanian'],
    'peru': ['peru', 'peruvian'],
    'philippines': ['philippines', 'philippine', 'filipino'],
    'poland': ['poland', 'polish'],
    'puerto rico': ['puerto rico'],
    
    # R
    'romania': ['romania', 'romanian'],
    'russia': ['russia', 'russian'],
    
    # S
    'saudi arabia': ['saudi arabia', 'saudi'],
    'senegal': ['senegal', 'senegalese'],
    'sierra leone': ['sierra leone'],
    'singapore': ['singapore'],
    'sint maarten': ['sint maarten', 'st maarten'],
    'slovakia': ['slovakia', 'slovak'],
    'south africa': ['south africa', 'south african'],
    'southern africa': ['southern africa'],
    'sri lanka': ['sri lanka', 'sri lankan'],
    'st kitts and nevis': ['st kitts and nevis', 'st. kitts & nevis', 'saint kitts and nevis'],
    'st lucia': ['st lucia', 'st. lucia', 'saint lucia'],
    'st vincent and the grenadines': ['st vincent and the grenadines', 'st. vincent & the grenadines', 'saint vincent and the grenadines'],
    'sub-saharan africa': ['sub-saharan africa', 'sub saharan africa'],
    'suriname': ['suriname'],
    'swaziland': ['swaziland', 'eswatini'],
    
    # T
    'tanzania': ['tanzania', 'tanzanian'],
    'thailand': ['thailand', 'thai'],
    'togo': ['togo'],
    'trinidad and tobago': ['trinidad and tobago', 'trinidad', 'tobago', 'tt'],
    'tunisia': ['tunisia', 'tunisian'],
    'turkey': ['turkey', 'turkish'],
    'turks and caicos': ['turks and caicos', 'turks & caicos', 'tci', 'turks & caicos islands'],
    
    # U
    'uganda': ['uganda', 'ugandan'],
    'united states': ['united states', 'usa', 'us', 'america', 'american'],
    'uruguay': ['uruguay', 'uruguayan'],
    
    # V
    'venezuela': ['venezuela', 'venezuelan'],
    'vietnam': ['vietnam', 'vietnamese'],
    
    # Y
    'yemen': ['yemen', 'yemeni'],
    
    # Z
    'zambia': ['zambia', 'zambian'],
    'zimbabwe': ['zimbabwe', 'zimbabwean']
}

# Enhanced region mapping with proper country associations based on World Bank regions
KNOWN_REGIONS = {
    'east asia and pacific': {
        'variations': ['east asia and pacific', 'east asia', 'pacific', 'eap', 'asia pacific'],
        'countries': ['china', 'indonesia', 'korea', 'laos', 'malaysia', 'mongolia', 'philippines', 
                     'thailand', 'vietnam', 'guam']
    },
    'europe and central asia': {
        'variations': ['europe and central asia', 'europe', 'central asia', 'eca', 'european'],
        'countries': ['austria', 'azerbaijan', 'bosnia herzegovina', 'bulgaria', 'czech republic', 
                     'estonia', 'germany', 'hungary', 'kyrgyz republic', 'lithuania', 'netherlands', 
                     'poland', 'romania', 'russia', 'slovakia', 'turkey']
    },
    'latin america and the caribbean': {
        'variations': ['latin america and the caribbean', 'latin america', 'caribbean', 'lac', 'latam'],
        'countries': ['argentina', 'aruba', 'bahamas', 'barbados', 'belize', 'bolivia', 'brazil', 
                     'chile', 'colombia', 'costa rica', 'curacao', 'dominica', 'dominican republic', 
                     'ecuador', 'el salvador', 'grenada', 'guatemala', 'guyana', 'haiti', 'honduras', 
                     'jamaica', 'mexico', 'panama', 'peru', 'puerto rico', 'sint maarten', 
                     'st kitts and nevis', 'st lucia', 'st vincent and the grenadines', 'suriname', 
                     'trinidad and tobago', 'turks and caicos', 'uruguay', 'venezuela']
    },
    'middle east and north africa': {
        'variations': ['middle east and north africa', 'middle east', 'north africa', 'mena', 'mea'],
        'countries': ['egypt', 'jordan', 'lebanon', 'morocco', 'oman', 'saudi arabia', 'tunisia', 'yemen']
    },
    'north america': {
        'variations': ['north america', 'na', 'north american'],
        'countries': ['canada', 'united states']
    },
    'south asia': {
        'variations': ['south asia', 'sa', 'south asian'],
        'countries': ['bangladesh', 'india', 'maldives', 'nepal', 'pakistan', 'sri lanka']
    },
    'sub saharan africa': {
        'variations': ['sub saharan africa', 'sub-saharan africa', 'africa', 'ssa', 'sub saharan', 'southern africa'],
        'countries': ['botswana', 'côte d\'ivoire', 'gabon', 'ghana', 'guinea', 'guinea-bissau', 
                     'kenya', 'liberia', 'madagascar', 'malawi', 'mauritania', 'mauritius', 
                     'namibia', 'nigeria', 'senegal', 'sierra leone', 'south africa', 'swaziland', 
                     'tanzania', 'togo', 'uganda', 'zambia', 'zimbabwe']
    },
    'global': {
        'variations': ['global', 'worldwide', 'international'],
        'countries': ['global']
    }
}

# Enhanced services with more variations and alternative wordings
KNOWN_SERVICES = {
    'due diligence': ['due diligence', 'dd', 'due-diligence', 'technical due diligence', 'commercial due diligence', 
                     'financial due diligence', 'environmental due diligence', 'diligence'],
    'feasibility study': ['feasibility study', 'feasibility', 'fs', 'feas study', 'feasibility analysis', 
                         'technical feasibility', 'economic feasibility', 'pre-feasibility'],
    "lender's engineer": ["lender's engineer", "lenders engineer", 'le', 'lender engineer', 'independent engineer',
                         'lenders technical advisor', 'lta'],
    "owner's engineer": ["owner's engineer", "owners engineer", 'oe', 'owner engineer', 'owners representative',
                        'project management', 'construction supervision'],
    'policy & regulation': ['policy & regulation', 'policy and regulation', 'policy', 'regulation', 'regulatory',
                           'policy analysis', 'regulatory framework', 'compliance'],
    'project development': ['project development', 'development', 'dev', 'project dev', 'business development',
                           'project planning', 'project structuring'],
    'transaction advisory': ['transaction advisory', 'transaction', 'advisory', 'ta', 'financial advisory',
                            'investment advisory', 'deal advisory', 'merger', 'acquisition']
}

# Enhanced sectors with STRICTER matching - only exact sector matches
KNOWN_SECTORS = {
    'renewable energy': ['renewable energy', 'renewable', 'renewables', 're', 'clean energy', 'green energy',
                        'sustainable energy', 'alternative energy'],
    'conventional energy': ['conventional energy', 'conventional', 'traditional energy', 'ce', 'fossil',
                           'fossil fuel', 'thermal power', 'conventional power'],
    'energy storage': ['energy storage', 'storage', 'battery storage', 'es', 'energy storage system',
                      'grid storage', 'utility storage'],
    'hydrogen': ['hydrogen', 'h2', 'green hydrogen', 'blue hydrogen', 'hydrogen energy', 'hydrogen power',
                'hydrogen production', 'hydrogen economy'],
    'lng to power': ['lng to power', 'lng power', 'lng-to-power', 'lng2power', 'lng'], 
    'other energy': ['other energy', 'misc energy', 'miscellaneous energy', 'mixed energy'],
    'water & wastewater': ['water & wastewater', 'water and wastewater', 'water', 'wastewater', 'ww',
                          'water treatment', 'sewage treatment', 'water infrastructure'],
    'other infrastructure': ['other infrastructure', 'infrastructure', 'other infra', 'infra',
                            'civil infrastructure', 'public infrastructure']
}

# Enhanced technologies with STRICTER matching - only exact technology matches
KNOWN_TECHNOLOGIES = {
    'wind': ['wind', 'wind power', 'wind energy', 'onshore wind', 'offshore wind', 'wind turbine',
            'wind farm', 'eolic'],
    'solar': ['solar', 'solar power', 'solar energy', 'pv', 'photovoltaic', 'solar pv', 'solar panel',
             'solar farm', 'solar plant', 'csp', 'concentrated solar power'],
    'hydro': ['hydro', 'hydropower', 'hydroelectric', 'hydro power', 'hydroelectric power',
             'small hydro', 'mini hydro', 'run of river'],
    'ulsd/diesel': ['ulsd/diesel', 'ulsd', 'diesel', 'ultra low sulfur diesel', 'diesel generator',
                   'diesel power', 'diesel engine'],
    'hfo': ['hfo', 'heavy fuel oil', 'fuel oil', 'residual fuel oil'],
    'others': ['others', 'other', 'misc', 'miscellaneous', 'mixed technology'],
    'nuclear': ['nuclear', 'nuclear power', 'nuclear energy', 'atomic power', 'nuclear plant'],
    'natural gas': ['natural gas', 'gas', 'ng', 'nat gas', 'gas turbine', 'gas engine', 'ccgt',
                   'combined cycle', 'gas power', 'liquefied natural gas'],
    'coal': ['coal', 'coal power', 'coal energy', 'coal plant', 'coal fired', 'thermal coal'],
    'green hydrogen': ['green hydrogen', 'green h2', 'renewable hydrogen', 'electrolytic hydrogen'],
    'geothermal': ['geothermal', 'geothermal energy', 'geothermal power', 'geothermal plant'],
    'bess': ['bess', 'battery energy storage system', 'battery storage', 'battery', 'lithium battery',
            'grid battery', 'utility battery'],
    'biomass': ['biomass', 'biomass energy', 'biomass power', 'bio energy', 'biofuel', 'biogas',
               'waste to energy', 'bagasse']
}

# ---- NEW: Client List Generation and Fuzzy Matching ----

def generate_client_list(projects: List[Dict]) -> List[str]:
    """Generate a comprehensive list of unique clients from the project data"""
    clients = set()
    
    for project in projects:
        # Get client from main field
        if project.get('client'):
            clients.add(project['client'].strip())
        
        # Get client from docx data
        if project.get('docx_data', {}).get('docx_client'):
            clients.add(project['docx_data']['docx_client'].strip())
    
    # Clean and filter clients
    cleaned_clients = []
    for client in clients:
        if client and len(client) > 2:  # Filter out very short or empty entries
            cleaned_clients.append(client)
    
    return sorted(list(set(cleaned_clients)))

def fuzzy_match_clients(query_term: str, client_list: List[str], threshold: float = 0.6) -> List[str]:
    """Fuzzy match client names using sequence matching"""
    query_lower = query_term.lower().strip()
    matches = []
    
    for client in client_list:
        client_lower = client.lower()
        
        # Exact match
        if query_lower == client_lower:
            matches.append(client)
            continue
        
        # Partial match
        if query_lower in client_lower or client_lower in query_lower:
            matches.append(client)
            continue
        
        # Fuzzy match using sequence matcher
        similarity = SequenceMatcher(None, query_lower, client_lower).ratio()
        if similarity >= threshold:
            matches.append(client)
            continue
        
        # Word-level fuzzy matching
        query_words = query_lower.split()
        client_words = client_lower.split()
        
        for query_word in query_words:
            if len(query_word) >= 3:  # Skip very short words
                for client_word in client_words:
                    if len(client_word) >= 3:
                        word_similarity = SequenceMatcher(None, query_word, client_word).ratio()
                        if word_similarity >= threshold:
                            matches.append(client)
                            break
                if client in matches:
                    break
    
    return list(set(matches))  # Remove duplicates

# ---- SPECIALIZED SEARCH FUNCTIONS ----
