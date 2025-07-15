import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import re

def generate_improved_word_document(uploaded_file, qualifications_text, relevant_projects):
    """
    Generate Word document with exact formatting preservation
    """
    try:
        # Load original document
        uploaded_file.seek(0)
        doc = docx.Document(uploaded_file)
        
        # Find Key Qualifications and Education sections in original
        key_qual_para = None
        education_para = None
        work_undertaken_para = None
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "Key Qualifications":
                key_qual_para = i
            elif "Education" in para.text:
                education_para = i
            elif "Work Undertaken that Best Illustrates Capability to Handle the Tasks Assigned" in para.text:
                work_undertaken_para = i
        
        # Remove existing content between Key Qualifications and Education
        if key_qual_para is not None and education_para is not None:
            paragraphs_to_remove = []
            for i in range(key_qual_para + 1, education_para):
                paragraphs_to_remove.append(doc.paragraphs[i])
            
            for para in paragraphs_to_remove:
                p = para._element
                p.getparent().remove(p)
        
        # Add new qualifications content after Key Qualifications heading
        if key_qual_para is not None:
            # Get the Key Qualifications paragraph to copy its formatting
            key_qual_element = doc.paragraphs[key_qual_para]._element
            parent = key_qual_element.getparent()
            
            # Split qualifications into paragraphs
            qual_paragraphs = qualifications_text.split('\n\n')
            
            # Insert each qualification paragraph right after Key Qualifications
            insert_position = parent.index(key_qual_element) + 1
            
            for qual_para in qual_paragraphs:
                if qual_para.strip():
                    # Create new paragraph element
                    new_p = docx.oxml.parser.parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    
                    # Insert at correct position
                    parent.insert(insert_position, new_p)
                    
                    # Create paragraph object
                    new_para = docx.text.paragraph.Paragraph(new_p, doc)
                    
                    # Add text with proper formatting
                    run = new_para.add_run(qual_para.strip())
                    run.font.name = 'Century Gothic'
                    run.font.size = Pt(11)
                    
                    # Apply US Body Text style if available
                    try:
                        new_para.style = doc.styles['US Body Text']
                    except:
                        pass
                    
                    insert_position += 1
        
        # Handle Work Undertaken section and projects table
        if work_undertaken_para is not None and relevant_projects:
            # Find the projects table (Table 5)
            table_to_replace = None
            if len(doc.tables) >= 5:
                table_to_replace = doc.tables[4]  # Table 5 (0-indexed)
            
            if table_to_replace:
                # Store original table formatting
                original_table_style = table_to_replace.style
                original_width = None
                try:
                    original_width = table_to_replace.columns[0].width
                except:
                    pass
                
                # Clear existing table content but preserve the first row (header)
                for row_idx, row in enumerate(table_to_replace.rows):
                    if row_idx == 0:
                        # Keep first row empty (dark blue header/divider)
                        for cell in row.cells:
                            cell.text = ""
                    else:
                        # Clear content rows
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.clear()
                
                # Adjust table rows to match project count (need +1 for header row)
                current_rows = len(table_to_replace.rows)
                needed_rows = len(relevant_projects) + 1  # +1 for header row
                
                # Remove extra rows (but keep at least 1 for header)
                if current_rows > needed_rows:
                    for _ in range(current_rows - needed_rows):
                        if len(table_to_replace.rows) > 1:
                            table_to_replace._element.remove(table_to_replace.rows[-1]._element)
                
                # Add missing rows
                elif needed_rows > current_rows:
                    for _ in range(needed_rows - current_rows):
                        new_row = table_to_replace.add_row()
                        # Set column width for new rows
                        if original_width:
                            try:
                                new_row.cells[0].width = original_width
                            except:
                                pass
                
                # Populate table with projects starting from row 1 (skip header row 0)
                for i, project in enumerate(relevant_projects):
                    row_index = i + 1  # Start from row 1, skip header row 0
                    if row_index < len(table_to_replace.rows):
                        cell = table_to_replace.rows[row_index].cells[0]
                        
                        # Clear existing content
                        cell.text = ""
                        
                        # Get the first paragraph in the cell
                        if cell.paragraphs:
                            paragraph = cell.paragraphs[0]
                        else:
                            paragraph = cell.add_paragraph()
                        
                        # Extract project title (text before first colon) and make it bold
                        if ':' in project:
                            # Split at first colon
                            colon_index = project.find(':')
                            title_part = project[:colon_index]
                            rest_part = project[colon_index:]  # Includes the colon and everything after
                            
                            # Add title part (bold)
                            title_run = paragraph.add_run(title_part)
                            title_run.bold = True
                            title_run.font.name = 'Century Gothic'
                            title_run.font.size = Pt(11)
                            
                            # Add rest of project description including colon (NOT bold)
                            rest_run = paragraph.add_run(rest_part)
                            rest_run.bold = False  # Explicitly set to not bold
                            rest_run.font.name = 'Century Gothic'
                            rest_run.font.size = Pt(11)
                        else:
                            # No colon found, treat entire text as title (bold)
                            title_run = paragraph.add_run(project)
                            title_run.bold = True
                            title_run.font.name = 'Century Gothic'
                            title_run.font.size = Pt(11)
                        
                        # Apply US Body Text style if available
                        try:
                            paragraph.style = doc.styles['US Body Text']
                        except:
                            pass
                        
                        # Preserve cell formatting
                        if original_width:
                            try:
                                cell.width = original_width
                            except:
                                pass
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        # Fallback: return original document if modification fails
        uploaded_file.seek(0)
        return uploaded_file.read() 